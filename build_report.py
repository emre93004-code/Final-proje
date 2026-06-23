# -*- coding: utf-8 -*-
import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches

# ==============================================================================
# --- 1. AYARLAR VE DİZİN OLUŞTURMA ---
# ==============================================================================
BASE = os.getcwd()
REP = os.path.join(BASE, "rapor")
os.makedirs(REP, exist_ok=True)


# Dosya yükleme yardımcısı
def load_file(filename):
    path = os.path.join(BASE, filename)
    if os.path.exists(path):
        print(f"[OK] {filename} dosyası başarıyla okundu.")
        return pd.read_csv(path)
    print(f"[HATA] {filename} dizinde bulunamadı!")
    return pd.DataFrame()


# Tablo oluşturma fonksiyonu
def df_to_table(document, df, title):
    document.add_heading(title, level=2)
    if df.empty: 
        document.add_paragraph("Bu tabloya ait veri kaynağı bulunamadı veya boş.")
        return
    
    # Word dokümanının okunabilirliğini korumak için çok uzun tablolarda ilk 15 satırı alalım
    df_preview = df.head(15)
    
    # Tabloyu oluşturma (Satır sayısı = başlık + veri satırları)
    t = document.add_table(rows=1, cols=len(df_preview.columns))
    t.style = "Light Grid Accent 1"
    
    # Başlıkları yazma
    hdr = t.rows[0].cells
    for j, c in enumerate(df_preview.columns):
        hdr[j].text = str(c)
        
    # Verileri yazma
    for _, row in df_preview.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df_preview.columns):
            cells[j].text = str(row[c])
            
    # Eğer veri seti kırpıldıysa altına küçük bir not ekleyelim
    if len(df) > 15:
        document.add_paragraph(f"* Not: Tablo çok uzun olduğu için yalnızca ilk 15 satır önizleme olarak listelenmiştir (Toplam satır: {len(df)}).")


# ==============================================================================
# --- 2. RAPORU İNŞA ETME BÖLÜMÜ ---
# ==============================================================================
# Doküman nesnesini oluşturuyoruz
doc = Document() 
doc.add_heading('Yapay Zeka Dersi - Ödev 2 Raporu', 0)

# --- Tablo 1: Proje Özeti ---
summary = load_file("summary.csv")
df_to_table(doc, summary, "Proje Özeti")

# --- Tablo 2: Kosinüs Benzerliği Değerlendirmesi ---
# Yüklediğin dosya adına uygun olarak güncellendi
cos = load_file("cosine_evaluation_summary.csv")
df_to_table(doc, cos, "Kosinüs Benzerliği Değerlendirmesi (Modellerin Genel Özeti)")

# --- Tablo 3: Modellerin İlk 5 Benzer Metin Çıktıları ---
# Yüklediğin dosya adına uygun olarak güncellendi
top5 = load_file("final_similarity_results (1).csv")
df_to_table(doc, top5, "Modellerin İlk 5 Benzer Metin Çıktıları")

# --- Tablo 4: Jaccard Matrisi (Ek Analiz Bölümü) ---
jaccard = load_file("jaccard_similarity_matrix.csv")
if not jaccard.empty:
    # İlk sütun isimsiz gelirse 'Model_Adı' olarak isimlendirelim
    if jaccard.columns[0] == "" or "Unnamed" in jaccard.columns[0]:
        jaccard.rename(columns={jaccard.columns[0]: "Model Adı"}, inplace=True)
    df_to_table(doc, jaccard, "Modeller Arası Jaccard Sıralama Kararlılığı Matrisi")


# ==============================================================================
# --- 3. TARTIŞMA VE SONUÇ BÖLÜMLERİ ---
# ==============================================================================
doc.add_heading("Tartışma ve Yorumlama", level=1)
doc.add_paragraph(
    "Modeller üzerinden yapılan analizlerde Lemmatized (kelime köklerine indirgenmiş) veri setleri "
    "ile eğitilen mimarilerin anlamsal tutarlılığı, Stemmed (ek kesme) yöntemine göre ezici bir "
    "üstünlük sağlamıştır. Stemming işleminin kelimelerin anlamsal bütünlüğünü bozduğu ve bağlam dışı "
    "sonuçlar ürettiği gözlemlenmiştir. CBOW mimarisi parametre değişimlerine karşı Skip-Gram'a kıyasla "
    "daha kararlı (robust) bir duruş sergilemiştir."
)

doc.add_heading("Sonuç ve Öneriler", level=1)
doc.add_paragraph(
    "Word2Vec modelleri insan kaynakları CV-İlan eşleştirme süreçlerinde hızlı ve pratik bir "
    "temel model (baseline) oluşturmaktadır. Sıfır Vektörü (Zero Vector) koruması sayesinde sistemin "
    "sözlük dışı (OOV) kelimelerde hata vermesi engellenmiştir. Gelecek çalışmalarda bağlamsal kelime "
    "temsilleri üretebilen Transformer tabanlı (BERT vb.) modellerin entegre edilmesi önerilmektedir."
)
